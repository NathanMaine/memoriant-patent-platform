"""Tests for core/export/docx_export.py."""
from __future__ import annotations

import io

import pytest
from docx import Document

from core.export.docx_export import DISCLAIMER, _add_paragraph, export_docx, export_dual
from core.models.application import DraftApplication, Embodiment, FilingFormat, Specification
from core.models.patent import Claim


# ---------------------------------------------------------------------------
# Helpers / fixtures
# ---------------------------------------------------------------------------

def _make_draft(
    filing_format: FilingFormat = FilingFormat.PROVISIONAL,
    abstract: str | None = "A system for optimizing widget performance using AI.",
    with_claims: bool = True,
    with_embodiments: bool = True,
    drawings_description: str | None = None,
) -> DraftApplication:
    embodiments = (
        [
            Embodiment(title="Cloud Implementation", description="Deployed as a cloud service."),
            Embodiment(title="Edge Implementation", description="Deployed on edge hardware."),
        ]
        if with_embodiments
        else []
    )
    claims = (
        [
            Claim(number=1, type="independent", text="A system comprising a processor."),
            Claim(number=2, type="dependent", text="The system of claim 1, further comprising memory.", depends_on=1),
        ]
        if with_claims
        else []
    )
    return DraftApplication(
        filing_format=filing_format,
        title="Smart Widget Optimization System",
        abstract=abstract,
        specification=Specification(
            background="Prior art lacked adaptive control.",
            summary="The present invention provides adaptive control.",
            detailed_description="The system comprises a processor and sensor array.",
            embodiments=embodiments,
        ),
        claims=claims,
        drawings_description=drawings_description,
    )


def _load_docx(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Basic output properties
# ---------------------------------------------------------------------------

def test_returns_bytes():
    """export_docx returns a bytes object."""
    draft = _make_draft()
    result = export_docx(draft)
    assert isinstance(result, bytes)


def test_returns_non_empty_bytes():
    """export_docx returns non-empty bytes."""
    draft = _make_draft()
    result = export_docx(draft)
    assert len(result) > 0


def test_starts_with_docx_magic_bytes():
    """DOCX is a zip file and must start with PK magic bytes."""
    draft = _make_draft()
    result = export_docx(draft)
    assert result[:2] == b"PK", "DOCX must start with PK zip header"


# ---------------------------------------------------------------------------
# Round-trip validation via python-docx
# ---------------------------------------------------------------------------

def test_round_trip_loads_with_python_docx():
    """Exported bytes can be loaded back into python-docx without error."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    assert doc is not None


def test_contains_title_text():
    """Exported DOCX contains the application title in ALL CAPS."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "SMART WIDGET OPTIMIZATION SYSTEM" in full_text


def test_contains_abstract_text():
    """Exported DOCX contains the abstract text."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "optimizing widget performance" in full_text


def test_contains_claim_text():
    """Exported DOCX contains numbered claim text."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "1. A system comprising a processor." in full_text
    assert "2. The system of claim 1" in full_text


def test_contains_specification_sections():
    """Exported DOCX contains all three specification section headings."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "BACKGROUND OF THE INVENTION" in full_text
    assert "SUMMARY OF THE INVENTION" in full_text
    assert "DETAILED DESCRIPTION OF THE INVENTION" in full_text


def test_contains_disclaimer_text():
    """Exported DOCX contains the filing disclaimer."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert DISCLAIMER in full_text


def test_contains_embodiment_subsections():
    """Exported DOCX contains embodiment subsection text."""
    draft = _make_draft()
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "Cloud Implementation" in full_text
    assert "Edge Implementation" in full_text


# ---------------------------------------------------------------------------
# Filing format variants
# ---------------------------------------------------------------------------

def test_works_with_provisional_format():
    """export_docx works with PROVISIONAL filing format."""
    draft = _make_draft(filing_format=FilingFormat.PROVISIONAL)
    result = export_docx(draft)
    assert result[:2] == b"PK"
    doc = _load_docx(result)
    assert "PROVISIONAL" in _all_text(doc).upper()


def test_works_with_nonprovisional_format():
    """export_docx works with NONPROVISIONAL filing format."""
    draft = _make_draft(filing_format=FilingFormat.NONPROVISIONAL)
    result = export_docx(draft)
    assert result[:2] == b"PK"
    doc = _load_docx(result)
    assert "NONPROVISIONAL" in _all_text(doc).upper()


def test_works_with_pct_format():
    """export_docx works with PCT filing format."""
    draft = _make_draft(filing_format=FilingFormat.PCT)
    result = export_docx(draft)
    assert result[:2] == b"PK"
    doc = _load_docx(result)
    assert "PCT" in _all_text(doc).upper()


# ---------------------------------------------------------------------------
# Optional field handling
# ---------------------------------------------------------------------------

def test_handles_missing_abstract():
    """export_docx works when abstract is None."""
    draft = _make_draft(abstract=None)
    result = export_docx(draft)
    assert result[:2] == b"PK"
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "No abstract provided" in full_text


def test_handles_missing_drawings_description():
    """export_docx works when drawings_description is None (no drawings section)."""
    draft = _make_draft(drawings_description=None)
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "BRIEF DESCRIPTION OF THE DRAWINGS" not in full_text


def test_handles_drawings_description_when_present():
    """export_docx includes drawings section when drawings_description is set."""
    draft = _make_draft(drawings_description="FIG. 1 shows the overall architecture.")
    result = export_docx(draft)
    doc = _load_docx(result)
    full_text = _all_text(doc)
    assert "BRIEF DESCRIPTION OF THE DRAWINGS" in full_text
    assert "FIG. 1 shows the overall architecture." in full_text


def test_handles_no_embodiments():
    """export_docx works when specification has no embodiments."""
    draft = _make_draft(with_embodiments=False)
    result = export_docx(draft)
    assert result[:2] == b"PK"
    doc = _load_docx(result)
    assert doc is not None


def test_handles_no_claims():
    """export_docx works when there are no claims."""
    draft = _make_draft(with_claims=False)
    result = export_docx(draft)
    assert result[:2] == b"PK"
    doc = _load_docx(result)
    assert doc is not None


def test_add_paragraph_with_alignment():
    """_add_paragraph alignment branch sets paragraph alignment when provided."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _add_paragraph(doc, "Centered text.", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Last paragraph should have CENTER alignment set
    last_para = doc.paragraphs[-1]
    assert last_para.text == "Centered text."
    assert last_para.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# export_dual
# ---------------------------------------------------------------------------

def test_export_dual_returns_tuple_of_two_bytes(monkeypatch):
    """export_dual returns a tuple of (docx_bytes, pdf_bytes)."""
    fake_pdf = b"%PDF-1.4 fake"
    monkeypatch.setattr("core.export.pdf_export.export_pdf", lambda draft: fake_pdf)
    draft = _make_draft()
    docx_bytes, pdf_bytes = export_dual(draft)
    assert isinstance(docx_bytes, bytes)
    assert isinstance(pdf_bytes, bytes)
    assert docx_bytes[:2] == b"PK"
    assert pdf_bytes == fake_pdf
