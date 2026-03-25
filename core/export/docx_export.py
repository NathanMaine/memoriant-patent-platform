"""DOCX export for patent draft applications — USPTO formatting."""
from __future__ import annotations

import io

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from core.models.application import DraftApplication, FilingFormat

logger = structlog.get_logger(__name__)

DISCLAIMER = (
    "This document was generated with AI assistance. "
    "Review by a qualified patent attorney is recommended before filing."
)


def _set_font(run, name: str = "Times New Roman", size: int = 12) -> None:
    run.font.name = name
    run.font.size = Pt(size)


def _add_paragraph(doc: Document, text: str, bold: bool = False, alignment=None) -> None:
    """Add a paragraph with Times New Roman 12pt, double-spaced."""
    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    _set_font(run)
    run.bold = bold
    _set_line_spacing(paragraph)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading paragraph (bold, Times New Roman 12pt)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_font(run)
    run.bold = True
    _set_line_spacing(paragraph)


def _set_line_spacing(paragraph) -> None:
    """Set double line spacing on a paragraph."""
    from docx.shared import Pt as DPt
    from docx.oxml.ns import qn as dqn
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "480")  # 240 = single, 480 = double
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


def export_docx(draft: DraftApplication) -> bytes:
    """Export a DraftApplication to a USPTO-formatted DOCX file.

    Returns the document as raw bytes.
    """
    log = logger.bind(draft_id=str(draft.id), filing_format=draft.filing_format)
    log.info("starting_docx_export")

    doc = Document()

    # -----------------------------------------------------------------------
    # Page setup: letter 8.5"x11", 1" top, 0.75" sides
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # -----------------------------------------------------------------------
    # Title page
    # -----------------------------------------------------------------------
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(draft.title.upper())
    _set_font(run)
    run.bold = True
    _set_line_spacing(title_paragraph)

    filing_paragraph = doc.add_paragraph()
    filing_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = filing_paragraph.add_run(f"Filing Format: {draft.filing_format.upper()}")
    _set_font(run)
    _set_line_spacing(filing_paragraph)

    # -----------------------------------------------------------------------
    # Abstract (separate page)
    # -----------------------------------------------------------------------
    _add_page_break(doc)
    _add_heading(doc, "ABSTRACT")

    if draft.abstract:
        _add_paragraph(doc, draft.abstract)
    else:
        _add_paragraph(doc, "(No abstract provided.)")

    # -----------------------------------------------------------------------
    # Specification sections (Background, Summary, Detailed Description)
    # -----------------------------------------------------------------------
    _add_page_break(doc)
    _add_heading(doc, "SPECIFICATION")

    _add_heading(doc, "BACKGROUND OF THE INVENTION")
    _add_paragraph(doc, draft.specification.background)

    _add_heading(doc, "SUMMARY OF THE INVENTION")
    _add_paragraph(doc, draft.specification.summary)

    _add_heading(doc, "DETAILED DESCRIPTION OF THE INVENTION")
    _add_paragraph(doc, draft.specification.detailed_description)

    # Embodiments as subsections
    for idx, embodiment in enumerate(draft.specification.embodiments, start=1):
        _add_heading(doc, f"Embodiment {idx}: {embodiment.title}")
        _add_paragraph(doc, embodiment.description)

    # Drawings description (optional)
    if draft.drawings_description:
        _add_heading(doc, "BRIEF DESCRIPTION OF THE DRAWINGS")
        _add_paragraph(doc, draft.drawings_description)

    # -----------------------------------------------------------------------
    # Claims (separate page)
    # -----------------------------------------------------------------------
    _add_page_break(doc)
    _add_heading(doc, "CLAIMS")

    for claim in draft.claims:
        claim_text = f"{claim.number}. {claim.text}"
        _add_paragraph(doc, claim_text)

    # -----------------------------------------------------------------------
    # Filing disclaimer
    # -----------------------------------------------------------------------
    disclaimer_paragraph = doc.add_paragraph()
    disclaimer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer_paragraph.add_run(DISCLAIMER)
    _set_font(run)
    run.italic = True
    _set_line_spacing(disclaimer_paragraph)

    # -----------------------------------------------------------------------
    # Serialize to bytes
    # -----------------------------------------------------------------------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    result = buffer.read()

    log.info("docx_export_complete", size_bytes=len(result))
    return result


def export_dual(draft: DraftApplication) -> tuple[bytes, bytes]:
    """Export a DraftApplication to both DOCX and PDF formats.

    Returns (docx_bytes, pdf_bytes).
    """
    from core.export.pdf_export import export_pdf

    log = logger.bind(draft_id=str(draft.id))
    log.info("starting_dual_export")

    docx_bytes = export_docx(draft)
    pdf_bytes = export_pdf(draft)

    log.info("dual_export_complete", docx_bytes=len(docx_bytes), pdf_bytes=len(pdf_bytes))
    return docx_bytes, pdf_bytes
